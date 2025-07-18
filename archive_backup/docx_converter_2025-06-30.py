"""
DOCX to Markdown converter for email attachments.

This module provides functionality to convert Microsoft Word documents (.docx)
to Markdown format with support for:
- Text extraction with style preservation
- Table conversion
- Image extraction
- Metadata extraction
- AI-ready text chunking
"""

import os
import json
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import logging

import mammoth
from bs4 import BeautifulSoup
import docx
from docx import Document

from email_parser.converters.base_converter import BaseConverter
from email_parser.exceptions.converter_exceptions import (
    ConversionError,
    ConfigurationError
)

logger = logging.getLogger(__name__)


class DocxConverter(BaseConverter):
    """Convert DOCX files to Markdown format with metadata extraction."""
    
    # Default configuration
    DEFAULT_CONFIG = {
        'max_file_size': 52428800,  # 50MB
        'output_format': 'both',  # json, html, both
        'extract_tables': True,
        'enable_chunking': False,  # Week 2 feature
        'max_chunk_tokens': 2000,
        'chunk_overlap': 200,
        'extract_metadata': True,
        'extract_styles': True,
        'include_comments': True,
        'extract_images': False,  # Week 2 feature
        'image_quality': 85,
        'max_image_size': 1200
    }
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize DOCX converter with configuration.
        
        Args:
            config: Configuration dictionary for the converter
            
        Raises:
            ConfigurationError: If initialization fails
        """
        # Merge with default config
        merged_config = self.DEFAULT_CONFIG.copy()
        if config:
            merged_config.update(config)
            
        super().__init__(merged_config)
        
        # Set configuration parameters from merged config
        self.extract_tables = self.config.get('extract_tables', True)
        self.extract_images = self.config.get('extract_images', False)
        self.extract_metadata = self.config.get('extract_metadata', True)
        self.enable_chunking = self.config.get('enable_chunking', False)
        
        logger.info(f"DOCX converter initialized with config: {self.config}")
    
    @property
    def supported_extensions(self) -> List[str]:
        """Return supported DOCX file extensions."""
        return ['.docx']
    
    @property  
    def supported_mime_types(self) -> List[str]:
        """Return supported DOCX MIME types."""
        return [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.ms-word.document.macroEnabled.12'
        ]
    
    @property
    def converter_name(self) -> str:
        """Return the converter name."""
        return "DOCX to Markdown Converter"
    
    def convert(self, input_path: Path, output_path: Optional[Path] = None) -> Path:
        """
        Convert DOCX file to Markdown.
        
        Args:
            input_path: Path to input DOCX file
            output_path: Optional path for output file (auto-generated if None)
            
        Returns:
            Path to the converted output file
            
        Raises:
            ConversionError: If conversion fails
        """
        try:
            # Validate the file
            self.validate_file(input_path)
            
            # Generate output path if not provided
            if output_path is None:
                output_path = self.generate_output_path(input_path, "docx")
            
            # Log conversion start
            self.log_conversion_start(input_path, output_path)
            start_time = logger.info.__self__.handlers[0].formatter.converter(logger.makeRecord(
                logger.name, logging.INFO, __file__, 0, "", (), None
            ).created) if logger.handlers else 0
            
            # Extract content using mammoth
            markdown_content = self._extract_with_mammoth(input_path)
            
            # Extract metadata if enabled
            metadata = {}
            if self.extract_metadata:
                metadata = self._extract_metadata(input_path)
            
            # Save markdown content
            with open(output_path, 'w', encoding='utf-8') as f:
                # Add metadata as YAML front matter if available
                if metadata:
                    f.write("---\n")
                    for key, value in metadata.items():
                        f.write(f"{key}: {value}\n")
                    f.write("---\n\n")
                f.write(markdown_content)
            
            # Save metadata as separate JSON file if enabled
            if metadata:
                metadata_path = output_path.with_suffix('.json')
                with open(metadata_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2, ensure_ascii=False)
            
            # Calculate duration (simplified for now)
            duration = 1.0  # Placeholder
            
            # Log successful conversion
            self.log_conversion_success(input_path, output_path, duration)
            
            return output_path
            
        except Exception as e:
            self.log_conversion_error(input_path, e)
            raise ConversionError(f"DOCX conversion failed: {str(e)}") from e
    
    def _extract_with_mammoth(self, file_path: Path) -> str:
        """
        Extract content from DOCX using mammoth.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted content as Markdown
        """
        try:
            # Basic mammoth extraction for Week 1
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                
            if result.messages:
                for message in result.messages:
                    logger.warning(f"Mammoth warning: {message}")
                    
            return result.value
            
        except Exception as e:
            raise ConversionError(f"Mammoth extraction failed: {str(e)}") from e
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing metadata
        """
        try:
            doc = Document(file_path)
            core_properties = doc.core_properties
            
            metadata = {
                'title': core_properties.title or '',
                'author': core_properties.author or '',
                'created': core_properties.created.isoformat() if core_properties.created else None,
                'modified': core_properties.modified.isoformat() if core_properties.modified else None,
                'subject': core_properties.subject or '',
                'keywords': core_properties.keywords or '',
                'category': core_properties.category or '',
                'comments': core_properties.comments or '',
                'revision': core_properties.revision,
                'word_count': self._estimate_word_count(doc),
                'paragraph_count': len(doc.paragraphs) if hasattr(doc, 'paragraphs') else 0,
                'table_count': len(doc.tables) if hasattr(doc, 'tables') else 0
            }
            
            # Filter out empty values
            return {k: v for k, v in metadata.items() if v}
            
        except Exception as e:
            logger.warning(f"Metadata extraction failed: {str(e)}")
            return {}
    
    def _estimate_word_count(self, doc: Document) -> int:
        """
        Estimate word count from document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Estimated word count
        """
        try:
            word_count = 0
            for paragraph in doc.paragraphs:
                if hasattr(paragraph, 'text'):
                    words = paragraph.text.split()
                    word_count += len(words)
            return word_count
        except:
            return 0