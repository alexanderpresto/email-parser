"""
DOCX to Markdown converter for email attachments.

This module provides functionality to convert Microsoft Word documents (.docx)
to Markdown format with support for:
- Text extraction with style preservation
- Table conversion
- Image extraction
- Metadata extraction
- AI-ready text chunking
"""

import os
import json
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import logging

import mammoth
from bs4 import BeautifulSoup
import docx
from docx import Document

from email_parser.converters.base_converter import BaseConverter
from email_parser.exceptions.converter_exceptions import (
    ConversionError,
    ConfigurationError
)

# Import Week 2 components
from email_parser.converters.docx.chunking import (
    create_chunker, ChunkingStrategy, DocumentChunk
)
from email_parser.converters.docx.metadata_extractor import (
    MetadataExtractor, PropertyAnalyzer
)
from email_parser.converters.docx.style_extractor import (
    StyleExtractor, StyleConverter, DocumentStyles
)
from email_parser.converters.docx.image_handler import (
    ImageHandler, ImageManifest
)

logger = logging.getLogger(__name__)


class DocxConverter(BaseConverter):
    """Convert DOCX files to Markdown format with metadata extraction."""
    
    # Default configuration
    DEFAULT_CONFIG = {
        'max_file_size': 52428800,  # 50MB
        'output_format': 'both',  # json, html, both
        'extract_tables': True,
        # AI Chunking settings
        'enable_chunking': True,  # Enable Week 2 feature
        'chunking_strategy': 'hybrid',  # token, semantic, hybrid
        'max_chunk_tokens': 2000,
        'chunk_overlap': 200,
        'preserve_structure': True,
        # Metadata settings
        'extract_metadata': True,
        'include_custom_properties': True,
        'analyze_metadata': True,
        # Style settings
        'extract_styles': True,
        'preserve_styles': True,
        'style_output_format': 'json',  # json, css
        'include_comments': True,
        # Image settings
        'extract_images': True,  # Enable Week 2 feature
        'image_quality': 85,
        'max_image_size': 1200,
        'generate_image_manifest': True
    }
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialize DOCX converter with configuration.
        
        Args:
            config: Configuration dictionary for the converter
            
        Raises:
            ConfigurationError: If initialization fails
        """
        # Merge with default config
        merged_config = self.DEFAULT_CONFIG.copy()
        if config:
            merged_config.update(config)
            
        super().__init__(merged_config)
        
        # Set configuration parameters from merged config
        self.extract_tables = self.config.get('extract_tables', True)
        self.extract_images = self.config.get('extract_images', True)
        self.extract_metadata = self.config.get('extract_metadata', True)
        self.extract_styles = self.config.get('extract_styles', True)
        self.enable_chunking = self.config.get('enable_chunking', True)
        
        # Initialize Week 2 components
        try:
            # Chunking
            self.chunker = None
            if self.enable_chunking:
                strategy = ChunkingStrategy(self.config.get('chunking_strategy', 'hybrid'))
                self.chunker = create_chunker(
                    strategy,
                    max_tokens=self.config.get('max_chunk_tokens', 2000),
                    overlap_tokens=self.config.get('chunk_overlap', 200)
                )
            
            # Metadata extractor
            self.metadata_extractor = MetadataExtractor() if self.extract_metadata else None
            
            # Style extractor
            self.style_extractor = StyleExtractor() if self.extract_styles else None
            
            # Image handler
            self.image_handler = None
            if self.extract_images:
                self.image_handler = ImageHandler(
                    extract_quality=self.config.get('image_quality', 85),
                    max_dimension=self.config.get('max_image_size', 1200)
                )
        except Exception as e:
            logger.warning(f"Could not initialize all Week 2 components: {e}")
            # Continue with basic functionality
        
        logger.info(f"DOCX converter initialized with Week 2 features: chunking={self.enable_chunking}, "
                   f"metadata={self.extract_metadata}, styles={self.extract_styles}, images={self.extract_images}")
    
    @property
    def supported_extensions(self) -> List[str]:
        """Return supported DOCX file extensions."""
        return ['.docx']
    
    @property  
    def supported_mime_types(self) -> List[str]:
        """Return supported DOCX MIME types."""
        return [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.ms-word.document.macroEnabled.12'
        ]
    
    @property
    def converter_name(self) -> str:
        """Return the converter name."""
        return "DOCX to Markdown Converter"
    
    def convert(self, input_path: Path, output_path: Optional[Path] = None) -> Path:
        """
        Convert DOCX file to Markdown with enhanced features.
        
        Args:
            input_path: Path to input DOCX file
            output_path: Optional path for output file (auto-generated if None)
            
        Returns:
            Path to the converted output file
            
        Raises:
            ConversionError: If conversion fails
        """
        try:
            # Validate the file
            self.validate_file(input_path)
            
            # Generate output path if not provided
            if output_path is None:
                output_path = self.generate_output_path(input_path, "docx")
            
            # Create output directory for additional files
            output_dir = output_path.parent / f"{output_path.stem}_docx_output"
            output_dir.mkdir(exist_ok=True)
            
            # Log conversion start
            self.log_conversion_start(input_path, output_path)
            
            # Extract content using mammoth
            markdown_content = self._extract_with_mammoth(input_path)
            
            # Week 2: Extract enhanced metadata
            metadata = {}
            metadata_analysis = {}
            if self.extract_metadata and self.metadata_extractor:
                try:
                    doc_metadata = self.metadata_extractor.extract(input_path)
                    metadata = doc_metadata.to_dict()
                    
                    # Analyze metadata if requested
                    if self.config.get('analyze_metadata', True):
                        metadata_analysis = PropertyAnalyzer.analyze_metadata(doc_metadata)
                except Exception as e:
                    logger.warning(f"Enhanced metadata extraction failed, using basic: {e}")
                    metadata = self._extract_metadata(input_path)
            
            # Week 2: Extract styles
            styles = None
            if self.extract_styles and self.style_extractor:
                try:
                    styles = self.style_extractor.extract_used_styles(str(input_path))
                    
                    # Save styles in requested format
                    style_format = self.config.get('style_output_format', 'json')
                    if style_format == 'css':
                        css_content = StyleConverter.to_css(styles)
                        css_path = output_dir / f"{output_path.stem}_styles.css"
                        css_path.write_text(css_content, encoding='utf-8')
                    elif style_format == 'json':
                        style_map = StyleConverter.to_style_map(styles)
                        styles_path = output_dir / f"{output_path.stem}_styles.json"
                        with open(styles_path, 'w', encoding='utf-8') as f:
                            json.dump(style_map, f, indent=2, ensure_ascii=False)
                except Exception as e:
                    logger.warning(f"Style extraction failed: {e}")
            
            # Week 2: Extract images
            image_manifest = {}
            if self.extract_images and self.image_handler:
                try:
                    images_dir = output_dir / "images"
                    saved_images = self.image_handler.extract_and_save_images(str(input_path), images_dir)
                    
                    if saved_images and self.config.get('generate_image_manifest', True):
                        # Create manifest
                        extracted_images = self.image_handler.extract_images(str(input_path))
                        image_manifest = ImageManifest.create_manifest(extracted_images)
                        ImageManifest.save_manifest(image_manifest, output_dir)
                        
                        logger.info(f"Extracted {len(saved_images)} images to {images_dir}")
                except Exception as e:
                    logger.warning(f"Image extraction failed: {e}")
            
            # Week 2: Apply chunking if enabled
            chunks = []
            if self.enable_chunking and self.chunker:
                try:
                    chunk_metadata = {
                        'document_title': metadata.get('title', ''),
                        'source_file': str(input_path)
                    }
                    chunks = self.chunker.chunk(markdown_content, chunk_metadata)
                    
                    # Save chunks as separate files
                    chunks_dir = output_dir / "chunks"
                    chunks_dir.mkdir(exist_ok=True)
                    
                    chunk_info = []
                    for chunk in chunks:
                        chunk_file = chunks_dir / f"chunk_{chunk.chunk_id:03d}.md"
                        with open(chunk_file, 'w', encoding='utf-8') as f:
                            # Add chunk metadata as YAML front matter
                            f.write("---\n")
                            f.write(f"chunk_id: {chunk.chunk_id}\n")
                            f.write(f"token_count: {chunk.token_count}\n")
                            f.write(f"start_index: {chunk.start_index}\n")
                            f.write(f"end_index: {chunk.end_index}\n")
                            if chunk.overlap_with_previous > 0:
                                f.write(f"overlap_with_previous: {chunk.overlap_with_previous}\n")
                            f.write("---\n\n")
                            f.write(chunk.content)
                        
                        chunk_info.append({
                            'chunk_id': chunk.chunk_id,
                            'file_path': str(chunk_file),
                            'token_count': chunk.token_count,
                            'overlap_with_previous': chunk.overlap_with_previous
                        })
                    
                    # Save chunk manifest
                    chunk_manifest = {
                        'total_chunks': len(chunks),
                        'chunking_strategy': self.config.get('chunking_strategy', 'hybrid'),
                        'max_tokens': self.config.get('max_chunk_tokens', 2000),
                        'overlap_tokens': self.config.get('chunk_overlap', 200),
                        'chunks': chunk_info
                    }
                    
                    chunk_manifest_path = output_dir / "chunk_manifest.json"
                    with open(chunk_manifest_path, 'w', encoding='utf-8') as f:
                        json.dump(chunk_manifest, f, indent=2, ensure_ascii=False)
                    
                    logger.info(f"Created {len(chunks)} chunks in {chunks_dir}")
                    
                except Exception as e:
                    logger.warning(f"Document chunking failed: {e}")
            
            # Save main markdown content
            with open(output_path, 'w', encoding='utf-8') as f:
                # Add metadata as YAML front matter if available
                if metadata:
                    f.write("---\n")
                    for key, value in metadata.items():
                        if value is not None and value != '':
                            f.write(f"{key}: {value}\n")
                    f.write("---\n\n")
                f.write(markdown_content)
            
            # Save comprehensive output manifest
            output_manifest = {
                'source_file': str(input_path),
                'conversion_time': logger.info.__self__.handlers[0].formatter.converter(logger.makeRecord(
                    logger.name, logging.INFO, __file__, 0, "", (), None
                ).created) if logger.handlers else None,
                'main_output': str(output_path),
                'output_directory': str(output_dir),
                'features_used': {
                    'enhanced_metadata': bool(self.metadata_extractor),
                    'style_extraction': bool(self.style_extractor),
                    'image_extraction': bool(self.image_handler),
                    'chunking': bool(self.chunker)
                },
                'metadata': metadata,
                'metadata_analysis': metadata_analysis,
                'chunks_count': len(chunks),
                'images_count': image_manifest.get('total_images', 0) if image_manifest else 0
            }
            
            manifest_path = output_dir / "conversion_manifest.json"
            with open(manifest_path, 'w', encoding='utf-8') as f:
                json.dump(output_manifest, f, indent=2, ensure_ascii=False, default=str)
            
            # Calculate duration (simplified for now)
            duration = 1.0  # Placeholder
            
            # Log successful conversion
            self.log_conversion_success(input_path, output_path, duration)
            logger.info(f"Enhanced DOCX conversion complete. Output directory: {output_dir}")
            
            return output_path
            
        except Exception as e:
            self.log_conversion_error(input_path, e)
            raise ConversionError(f"DOCX conversion failed: {str(e)}") from e
    
    def _extract_with_mammoth(self, file_path: Path) -> str:
        """
        Extract content from DOCX using mammoth.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted content as Markdown
        """
        try:
            # Basic mammoth extraction for Week 1
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                
            if result.messages:
                for message in result.messages:
                    logger.warning(f"Mammoth warning: {message}")
                    
            return result.value
            
        except Exception as e:
            raise ConversionError(f"Mammoth extraction failed: {str(e)}") from e
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing metadata
        """
        try:
            doc = Document(file_path)
            core_properties = doc.core_properties
            
            metadata = {
                'title': core_properties.title or '',
                'author': core_properties.author or '',
                'created': core_properties.created.isoformat() if core_properties.created else None,
                'modified': core_properties.modified.isoformat() if core_properties.modified else None,
                'subject': core_properties.subject or '',
                'keywords': core_properties.keywords or '',
                'category': core_properties.category or '',
                'comments': core_properties.comments or '',
                'revision': core_properties.revision,
                'word_count': self._estimate_word_count(doc),
                'paragraph_count': len(doc.paragraphs) if hasattr(doc, 'paragraphs') else 0,
                'table_count': len(doc.tables) if hasattr(doc, 'tables') else 0
            }
            
            # Filter out empty values
            return {k: v for k, v in metadata.items() if v}
            
        except Exception as e:
            logger.warning(f"Metadata extraction failed: {str(e)}")
            return {}
    
    def _estimate_word_count(self, doc: Document) -> int:
        """
        Estimate word count from document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Estimated word count
        """
        try:
            word_count = 0
            for paragraph in doc.paragraphs:
                if hasattr(paragraph, 'text'):
                    words = paragraph.text.split()
                    word_count += len(words)
            return word_count
        except:
            return 0
    
    def convert_standalone(self, file_path: Path, output_dir: Path, 
                          options: Optional[Dict[str, Any]] = None) -> Path:
        """
        Convert a DOCX file standalone without email context.
        
        Args:
            file_path: Path to the DOCX file to convert
            output_dir: Directory where output should be saved
            options: Optional conversion options
            
        Returns:
            Path to the converted file
            
        Raises:
            ConversionError: If conversion fails
        """
        # Generate output filename
        output_filename = f"{file_path.stem}.md"
        output_path = output_dir / output_filename
        
        # Use existing convert method with the specified output path
        return self.convert(file_path, output_path)